import asyncio
import os
import io
import re
import random
import string
import hashlib
import json
from datetime import datetime, timedelta
from typing import Optional, List
import base64

from dotenv import load_dotenv
load_dotenv()

from aiogram import Bot, Dispatcher, Router, F
from aiogram.types import (
    Message, CallbackQuery, BufferedInputFile,
    InlineKeyboardMarkup, InlineKeyboardButton, ReplyKeyboardMarkup,
    KeyboardButton, ReplyKeyboardRemove
)
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage

from docx import Document
from docx.oxml import parse_xml
from motor.motor_asyncio import AsyncIOMotorClient
from bson import ObjectId

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm

from PIL import Image as PILImage

# ================= CONFIG =================
BOT_TOKEN = os.getenv("BOT_TOKEN")
MONGO_URI = os.getenv("MONGO_URI", "mongodb://localhost:27017")
DB_NAME = os.getenv("DB_NAME", "fizika_test_bot")
ADMIN_IDS = [int(x.strip()) for x in os.getenv("ADMIN_IDS", "").split(",") if x.strip()]

DEFAULT_QUESTION_COUNT = int(os.getenv("DEFAULT_QUESTION_COUNT", 10))
DEFAULT_TIME_LIMIT = int(os.getenv("DEFAULT_TIME_LIMIT", 30))
PIN_EXPIRY_DAYS = int(os.getenv("PIN_EXPIRY_DAYS", 7))
MAX_IMAGE_SIZE = int(os.getenv("MAX_IMAGE_SIZE", 50000))
MAX_IMAGE_DIMENSION = int(os.getenv("MAX_IMAGE_DIMENSION", 800))

if not BOT_TOKEN:
    raise ValueError("❌ BOT_TOKEN topilmadi!")

# ================= DATABASE =================
client = AsyncIOMotorClient(
    MONGO_URI,
    serverSelectionTimeoutMS=30000,
    connectTimeoutMS=30000,
    socketTimeoutMS=30000,
    retryWrites=True,
    w="majority"
)
db = client[DB_NAME]

questions_col = db.questions
results_col = db.results
pins_col = db.pins
users_col = db.users
images_col = db.images
pin_batches_col = db.pin_batches

# ================= STATES =================
class AdminStates(StatesGroup):
    waiting_word = State()
    waiting_topic = State()
    waiting_grade = State()
    waiting_difficulty = State()
    add_manual_question = State()
    add_manual_topic = State()
    add_manual_options = State()
    add_manual_answer = State()
    pin_count = State()
    pin_config = State()
    pin_reset_select = State()
    delete_confirm = State()  # YANGI
    delete_by_grade = State()  # YANGI
    delete_by_topic = State()  # YANGI

class StudentStates(StatesGroup):
    waiting_pin = State()
    waiting_name = State()
    taking_test = State()
    waiting_text_answer = State()

class TeacherStates(StatesGroup):
    waiting_pin_for_report = State()
    selecting_report_type = State()  # YANGI
# ================= HELPERS =================
def generate_pin() -> str:
    return ''.join(random.choices(string.digits, k=8))

def compress_image(image_data: bytes) -> bytes:
    try:
        img = PILImage.open(io.BytesIO(image_data))
        if max(img.size) > MAX_IMAGE_DIMENSION:
            ratio = MAX_IMAGE_DIMENSION / max(img.size)
            new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
            img = img.resize(new_size, PILImage.LANCZOS)
        if img.mode in ('RGBA', 'P'):
            img = img.convert('RGB')
        output = io.BytesIO()
        quality = 85
        img.save(output, format='JPEG', quality=quality, optimize=True)
        while output.tell() > MAX_IMAGE_SIZE and quality > 20:
            output = io.BytesIO()
            quality -= 10
            img.save(output, format='JPEG', quality=quality, optimize=True)
        return output.getvalue()
    except Exception as e:
        print(f"Image compression error: {e}")
        return image_data

async def save_image(image_data: bytes) -> str:
    try:
        compressed = compress_image(image_data)
        img_hash = hashlib.md5(compressed).hexdigest()
        existing = await images_col.find_one({"hash": img_hash})
        if existing:
            return str(existing["_id"])
        result = await images_col.insert_one({
            "hash": img_hash,
            "data": base64.b64encode(compressed).decode(),
            "size": len(compressed)
        })
        return str(result.inserted_id)
    except Exception as e:
        print(f"Save image error: {e}")
        return None

async def get_image(image_id: str) -> Optional[bytes]:
    try:
        img = await images_col.find_one({"_id": ObjectId(image_id)})
        if img:
            return base64.b64decode(img["data"])
    except Exception as e:
        print(f"Get image error: {e}")
    return None

async def parse_word_with_images(file_path: str) -> List[dict]:
    """Word hujjatni rasmlar bilan parse qilish - YANGI VERSIYA"""
    doc = Document(file_path)
    questions = []
    current_q = None
    
    # Barcha rasmlarni paragraf indeksi bilan olish
    para_images = {}  # {para_index: [image_data, ...]}
    
    for para_idx, para in enumerate(doc.paragraphs):
        images = []
        try:
            for run in para.runs:
                # Inline rasmlar
                for drawing in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                    blip = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    if blip is not None:
                        embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed and embed in doc.part.rels:
                            try:
                                image_data = doc.part.rels[embed].target_part.blob
                                images.append(image_data)
                            except:
                                pass
                
                # Anchor rasmlar
                for anchor in run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor'):
                    blip = anchor.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    if blip is not None:
                        embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed and embed in doc.part.rels:
                            try:
                                image_data = doc.part.rels[embed].target_part.blob
                                images.append(image_data)
                            except:
                                pass
        except Exception as e:
            print(f"Para {para_idx} rasm xatosi: {e}")
        
        if images:
            para_images[para_idx] = images
            print(f"✓ Para {para_idx}: {len(images)} ta rasm")
    
    print(f"Jami {sum(len(imgs) for imgs in para_images.values())} ta rasm topildi")
    
    # Paragraflarni qayta ishlash
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not text:
            continue
        
        # Yangi savol
        if re.match(r'^\d+[\.\)]\s*', text):
            if current_q and current_q.get('text'):
                questions.append(current_q)
            
            q_text = re.sub(r'^\d+[\.\)]\s*', '', text)
            current_q = {
                'text': q_text,
                'options': [],
                'answer': None,
                'images': [],
                'type': 'choice',
                'explanation': '',
                'para_index': para_idx
            }
            
            # Shu paragrafda rasm bormi
            if para_idx in para_images:
                for img_data in para_images[para_idx]:
                    img_id = await save_image(img_data)
                    if img_id:
                        current_q['images'].append(img_id)
                        print(f"  → Savol {len(questions)+1} ga rasm: {img_id}")
            
            # Keyingi paragrafda rasm bormi tekshirish (savol keyin rasm)
            next_idx = para_idx + 1
            if next_idx in para_images:
                # Keyingi paragraf matn bo'sh yoki variant bo'lsa, rasm savolga tegishli
                next_para = doc.paragraphs[next_idx] if next_idx < len(doc.paragraphs) else None
                if next_para:
                    next_text = next_para.text.strip()
                    # Agar keyingi qator variant emas yoki bo'sh bo'lsa, rasm savolga tegishli
                    if not re.match(r'^[A-Da-d][\.\)]\s*', next_text):
                        for img_data in para_images[next_idx]:
                            img_id = await save_image(img_data)
                            if img_id and img_id not in current_q['images']:
                                current_q['images'].append(img_id)
                                print(f"  → Savol {len(questions)+1} ga rasm (keyingi): {img_id}")
        
        # Variant
        elif re.match(r'^[A-Da-d][\.\)]\s*', text) and current_q:
            opt_text = re.sub(r'^[A-Da-d][\.\)]\s*', '', text)
            current_q['options'].append(opt_text)
        
        # Javob
        elif text.lower().startswith('javob:') and current_q:
            answer = text.split(':', 1)[1].strip()
            if answer.upper() in ['A', 'B', 'C', 'D']:
                current_q['answer'] = ord(answer.upper()) - ord('A')
            else:
                current_q['answer'] = answer
                current_q['type'] = 'text'
        
        # Tushuntirish
        elif text.lower().startswith("tushuntirish:") and current_q:
            current_q['explanation'] = text.split(':', 1)[1].strip()
        
        # Davomi
        elif current_q and text:
            current_q['text'] += ' ' + text
    
    # Oxirgi savol
    if current_q and current_q.get('text'):
        questions.append(current_q)
    
    # para_index ni o'chirish
    for q in questions:
        q.pop('para_index', None)
    
    print(f"\n✅ {len(questions)} ta savol parse qilindi")
    for i, q in enumerate(questions, 1):
        print(f"  {i}. {q['text'][:50]}... | {len(q.get('images', []))} rasm | {len(q.get('options', []))} variant")
    
    return questions

def generate_pins_pdf(pins_data: List[dict], batch_info: dict) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
    elements = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, alignment=1, spaceAfter=20)
    elements.append(Paragraph("🔑 PIN KODLAR", title_style))
    elements.append(Spacer(1, 0.3*cm))
    
    info_style = ParagraphStyle('Info', parent=styles['Normal'], fontSize=10, spaceAfter=6)
    elements.append(Paragraph(f"<b>Sinf:</b> {batch_info['grade']}", info_style))
    elements.append(Paragraph(f"<b>Mavzu:</b> {batch_info['topic']}", info_style))
    elements.append(Paragraph(f"<b>Savollar:</b> {batch_info['question_count']} ta", info_style))
    elements.append(Paragraph(f"<b>Vaqt:</b> {batch_info['time_limit']} daqiqa", info_style))
    
    multi_use = batch_info.get('multi_use', False)
    max_attempts = batch_info.get('max_attempts', 1)
    attempts_text = "Cheksiz" if (multi_use and max_attempts >= 999) else f"{max_attempts} marta"
    elements.append(Paragraph(f"<b>Urinishlar:</b> {attempts_text}", info_style))
    elements.append(Paragraph(f"<b>Yaratildi:</b> {datetime.now().strftime('%d.%m.%Y %H:%M')}", info_style))
    elements.append(Spacer(1, 0.5*cm))
    
    table_data = [["№", "PIN KOD", "O'QUVCHI", "HOLAT"]]
    for i, pin in enumerate(pins_data, 1):
        table_data.append([str(i), pin['pin'], "", "Faol"])
    
    table = Table(table_data, colWidths=[1.5*cm, 4*cm, 8*cm, 3*cm])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4CAF50')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(table)
    
    doc.build(elements)
    buffer.seek(0)
    return buffer.read()

def generate_pins_json(pins_data: List[dict], batch_info: dict) -> str:
    export_data = {
        "batch_info": {
            "grade": batch_info['grade'],
            "topic": batch_info['topic'],
            "question_count": batch_info['question_count'],
            "time_limit": batch_info['time_limit'],
            "multi_use": batch_info.get('multi_use', False),
            "max_attempts": batch_info.get('max_attempts', 1),
            "created_at": datetime.now().isoformat(),
            "total_pins": len(pins_data)
        },
        "pins": [{"number": i, "pin": pin['pin'], "student": "", "status": "active"} 
                 for i, pin in enumerate(pins_data, 1)]
    }
    return json.dumps(export_data, ensure_ascii=False, indent=2)

# ================= BOT =================
bot = Bot(token=BOT_TOKEN)
dp = Dispatcher(storage=MemoryStorage())
router = Router()

# ================= KEYBOARDS =================
def admin_menu():
    return ReplyKeyboardMarkup(keyboard=[
        [KeyboardButton(text="📤 Word yuklash"), KeyboardButton(text="➕ Savol qo'shish")],
        [KeyboardButton(text="🔑 PIN yaratish"), KeyboardButton(text="📋 PIN boshqaruv")],
        [KeyboardButton(text="📊 Natijalar"), KeyboardButton(text="📈 Statistika")],
        [KeyboardButton(text="🗑 Savollarni o'chirish"), KeyboardButton(text="⚙️ Sozlamalar")]
    ], resize_keyboard=True)

def delete_questions_kb():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🗑 Sinf bo'yicha", callback_data="delq_by_grade")],
        [InlineKeyboardButton(text="🗑 Mavzu bo'yicha", callback_data="delq_by_topic")],
        [InlineKeyboardButton(text="🗑 Alohida savollar", callback_data="delq_individual")],
        [InlineKeyboardButton(text="🗑 Barchasini o'chirish", callback_data="delq_all")],
        [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="cancel")]
    ])

def report_type_kb():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📄 Qisqacha hisobot", callback_data="report_summary")],
        [InlineKeyboardButton(text="📋 Batafsil hisobot", callback_data="report_detailed")],
        [InlineKeyboardButton(text="📊 Ikkala hisobot", callback_data="report_both")]
    ])

def student_menu():
    return ReplyKeyboardMarkup(keyboard=[
        [KeyboardButton(text="📝 Test boshlash")],
        [KeyboardButton(text="📊 Natijalarim")]
    ], resize_keyboard=True)

def grade_kb():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="7-sinf", callback_data="grade_7")],
        [InlineKeyboardButton(text="8-sinf", callback_data="grade_8")],
        [InlineKeyboardButton(text="9-sinf", callback_data="grade_9")]
    ])

def diff_kb():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📗 Bilish", callback_data="diff_bilish")],
        [InlineKeyboardButton(text="📙 Qo'llash", callback_data="diff_qollash")],
        [InlineKeyboardButton(text="📕 Mulohaza", callback_data="diff_mulohaza")],
        [InlineKeyboardButton(text="🔀 Aralash", callback_data="diff_aralash")]
    ])

def pin_settings_kb():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="1️⃣ Bir martalik", callback_data="pinset_single")],
        [InlineKeyboardButton(text="♾️ Ko'p martalik", callback_data="pinset_multi")],
        [InlineKeyboardButton(text="🔢 Urinishlar soni", callback_data="pinset_attempts")]
    ])

def pin_management_kb():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📋 Barcha PIN'lar", callback_data="pinmgmt_list")],
        [InlineKeyboardButton(text="🔄 Reset", callback_data="pinmgmt_reset")],
        [InlineKeyboardButton(text="📊 Statistika", callback_data="pinmgmt_stats")]
    ])

def navigation_kb(current: int, total: int, answers: dict, idx: int):
    """Navigatsiya klaviaturasi - YANGILANGAN"""
    buttons = []
    
    # Savollar ro'yxati (3 qatorda 5 tadan)
    rows = []
    for i in range(0, total, 5):
        row = []
        for j in range(i, min(i + 5, total)):
            # Javob berilganmi tekshirish
            if str(j) in answers:
                emoji = "✅"  # Javob berilgan
            else:
                emoji = "⬜"  # Javob berilmagan
            
            # Hozirgi savol bo'lsa
            if j == current:
                text = f"➡️ {j+1}"
            else:
                text = f"{emoji} {j+1}"
            
            row.append(InlineKeyboardButton(
                text=text,
                callback_data=f"goto_{idx}_{j}"
            ))
        rows.append(row)
    
    buttons.extend(rows)
    
    # Navigatsiya tugmalari
    nav_row = []
    if current > 0:
        nav_row.append(InlineKeyboardButton(text="⬅️ Oldingi", callback_data=f"nav_{idx}_prev"))
    
    nav_row.append(InlineKeyboardButton(
        text=f"📊 {current+1}/{total}",
        callback_data="nav_info"
    ))
    
    if current < total - 1:
        nav_row.append(InlineKeyboardButton(text="Keyingi ➡️", callback_data=f"nav_{idx}_next"))
    
    buttons.append(nav_row)
    
    # Tugallash tugmasi
    buttons.append([InlineKeyboardButton(
        text="✅ Testni yakunlash",
        callback_data=f"finish_{idx}"
    )])
    
    return InlineKeyboardMarkup(inline_keyboard=buttons)

def ans_kb(options: List[str], idx: int, current: int):
    """Javob variantlari klaviaturasi"""
    buttons = []
    for i, opt in enumerate(options):
        letter = chr(65 + i)
        text = f"{letter}) {opt[:40]}{'...' if len(opt) > 40 else ''}"
        buttons.append([InlineKeyboardButton(
            text=text,
            callback_data=f"ans_{idx}_{current}_{i}"
        )])
    return InlineKeyboardMarkup(inline_keyboard=buttons)

# ================= HANDLERS =================
@router.message(Command("start"))
async def cmd_start(msg: Message):
    if msg.from_user.id in ADMIN_IDS:
        await msg.answer("👋 Salom Admin!\n\n📚 Fizika Test Bot", reply_markup=admin_menu())
    else:
        await users_col.update_one(
            {"user_id": msg.from_user.id},
            {"$set": {"user_id": msg.from_user.id, "name": msg.from_user.full_name, "reg": datetime.now()}},
            upsert=True
        )
        await msg.answer("👋 Fizika Test Botiga xush kelibsiz!", reply_markup=student_menu())

# ===== WORD UPLOAD =====
@router.message(F.text == "📤 Word yuklash")
async def upload_word(msg: Message, state: FSMContext):
    if msg.from_user.id not in ADMIN_IDS: return
    await msg.answer(
        "📄 Word faylni yuboring\n\n"
        "Format:\n"
        "1. Savol matni (rasm bo'lsa shu yerda)\n"
        "A) Variant 1\n"
        "B) Variant 2\n"
        "C) Variant 3\n"
        "D) Variant 4\n"
        "Javob: A"
    )
    await state.set_state(AdminStates.waiting_word)

@router.message(AdminStates.waiting_word, F.document)
async def process_word(msg: Message, state: FSMContext):
    if not msg.document.file_name.endswith('.docx'):
        await msg.answer("❌ Faqat .docx!"); return
    
    status = await msg.answer("⏳ Yuklanmoqda va rasmlar qayta ishlanmoqda...")
    file = await bot.get_file(msg.document.file_id)
    path = f"temp_{msg.from_user.id}.docx"
    await bot.download_file(file.file_path, path)
    
    try:
        questions = await parse_word_with_images(path)
        if not questions:
            await status.edit_text("❌ Savollar topilmadi!"); return
        
        images_count = sum(len(q.get('images', [])) for q in questions)
        await state.update_data(questions=questions)
        await status.edit_text(
            f"✅ {len(questions)} ta savol topildi!\n"
            f"🖼 {images_count} ta rasm yuklandi\n\n"
            "Sinf tanlang:",
            reply_markup=grade_kb()
        )
        await state.set_state(AdminStates.waiting_grade)
    except Exception as e:
        await status.edit_text(f"❌ Xato: {e}")
        print(f"Word parse error: {e}")
    finally:
        if os.path.exists(path): os.remove(path)

@router.callback_query(AdminStates.waiting_grade, F.data.startswith("grade_"))
async def word_grade(cb: CallbackQuery, state: FSMContext):
    grade = int(cb.data.split("_")[1])
    await state.update_data(grade=grade)
    await cb.message.edit_text(f"📚 {grade}-sinf\n\nMavzu nomini yozing:")
    await state.set_state(AdminStates.waiting_topic)

@router.message(AdminStates.waiting_topic)
async def word_topic(msg: Message, state: FSMContext):
    await state.update_data(topic=msg.text.strip())
    await msg.answer("Qiyinlik darajasini tanlang:", reply_markup=diff_kb())
    await state.set_state(AdminStates.waiting_difficulty)

@router.callback_query(AdminStates.waiting_difficulty, F.data.startswith("diff_"))
async def word_diff(cb: CallbackQuery, state: FSMContext):
    diff_map = {"bilish": "Bilish", "qollash": "Qo'llash", "mulohaza": "Mulohaza", "aralash": "Aralash"}
    diff = diff_map.get(cb.data.split("_")[1], "Aralash")
    data = await state.get_data()
    
    status = await cb.message.edit_text("⏳ Savollar saqlanmoqda...")
    
    for q in data['questions']:
        q.update({
            'grade': data['grade'],
            'topic': data['topic'],
            'difficulty': diff if diff != "Aralash" else random.choice(["Bilish", "Qo'llash", "Mulohaza"]),
            'created_at': datetime.now(),
            'created_by': cb.from_user.id
        })
        await questions_col.insert_one(q)
    
    await status.edit_text(f"✅ {len(data['questions'])} savol saqlandi!")
    await state.clear()

# ===== MANUAL ADD =====
@router.message(F.text == "➕ Savol qo'shish")
async def add_q(msg: Message, state: FSMContext):
    if msg.from_user.id not in ADMIN_IDS: return
    await msg.answer("Sinf:", reply_markup=grade_kb())
    await state.set_state(AdminStates.add_manual_question)

@router.callback_query(AdminStates.add_manual_question, F.data.startswith("grade_"))
async def add_q_grade(cb: CallbackQuery, state: FSMContext):
    await state.update_data(grade=int(cb.data.split("_")[1]))
    await cb.message.edit_text("Mavzu nomini kiriting:")
    await state.set_state(AdminStates.add_manual_topic)

@router.message(AdminStates.add_manual_topic)
async def add_q_topic(msg: Message, state: FSMContext):
    await state.update_data(topic=msg.text.strip())
    await msg.answer("Savol matnini yozing (rasm bo'lsa avval rasmni yuboring):")
    await state.set_state(AdminStates.add_manual_options)

@router.message(AdminStates.add_manual_options)
async def add_q_text(msg: Message, state: FSMContext):
    if msg.photo:
        photo = msg.photo[-1]
        file = await bot.get_file(photo.file_id)
        file_data = await bot.download_file(file.file_path)
        img_id = await save_image(file_data.read())
        
        if img_id:
            data = await state.get_data()
            await state.update_data(images=data.get('images', []) + [img_id])
            await msg.answer("✅ Rasm qo'shildi! Yana rasm yoki savol matnini yuboring:")
        else:
            await msg.answer("❌ Rasmni saqlashda xatolik!")
        return
    
    await state.update_data(q_text=msg.text.strip())
    await msg.answer("Variantlarni kiriting (har birini yangi qatordan):\n\nMasalan:\n5 kg\n10 kg\n15 kg\n20 kg")
    await state.set_state(AdminStates.add_manual_answer)

@router.message(AdminStates.add_manual_answer)
async def add_q_opts(msg: Message, state: FSMContext):
    data = await state.get_data()
    
    if 'options' not in data:
        opts = [o.strip() for o in msg.text.split('\n') if o.strip()]
        if len(opts) < 2:
            await msg.answer("❌ Kamida 2 ta variant!")
            return
        
        await state.update_data(options=opts)
        btns = [[InlineKeyboardButton(text=f"{chr(65+i)}) {o}", callback_data=f"correct_{i}")] for i, o in enumerate(opts)]
        await msg.answer("To'g'ri javobni tanlang:", reply_markup=InlineKeyboardMarkup(inline_keyboard=btns))

@router.callback_query(F.data.startswith("correct_"))
async def add_q_correct(cb: CallbackQuery, state: FSMContext):
    ans = int(cb.data.split("_")[1])
    data = await state.get_data()
    
    await questions_col.insert_one({
        'text': data['q_text'],
        'options': data['options'],
        'answer': ans,
        'images': data.get('images', []),
        'type': 'choice',
        'explanation': '',
        'grade': data['grade'],
        'topic': data['topic'],
        'difficulty': 'Bilish',
        'created_at': datetime.now(),
        'created_by': cb.from_user.id
    })
    
    await cb.message.edit_text(f"✅ Savol saqlandi!\n🖼 {len(data.get('images', []))} ta rasm")
    await state.clear()

# ===== PIN CREATION =====
@router.message(F.text == "🔑 PIN yaratish")
async def create_pin(msg: Message, state: FSMContext):
    if msg.from_user.id not in ADMIN_IDS: return
    await msg.answer("Sinf:", reply_markup=grade_kb())
    await state.update_data(action="create_pin")

@router.callback_query(F.data.startswith("grade_"))
async def pin_grade(cb: CallbackQuery, state: FSMContext):
    data = await state.get_data()
    if data.get('action') != 'create_pin': return
    
    grade = int(cb.data.split("_")[1])
    topics = await questions_col.distinct("topic", {"grade": grade})
    
    if not topics:
        await cb.message.edit_text("❌ Savollar yo'q!")
        return
    
    await state.update_data(grade=grade)
    btns = [[InlineKeyboardButton(text=t, callback_data=f"pintopic_{t}")] for t in topics]
    await cb.message.edit_text("Mavzu:", reply_markup=InlineKeyboardMarkup(inline_keyboard=btns))

@router.callback_query(F.data.startswith("pintopic_"))
async def pin_topic(cb: CallbackQuery, state: FSMContext):
    topic = cb.data.replace("pintopic_", "")
    await state.update_data(topic=topic)
    await cb.message.edit_text(f"📝 {topic}\n\nNechta PIN? (1-100):")
    await state.set_state(AdminStates.pin_count)

@router.message(AdminStates.pin_count)
async def pin_count(msg: Message, state: FSMContext):
    try:
        count = int(msg.text.strip())
        if count < 1 or count > 100:
            await msg.answer("❌ 1-100!")
            return
    except:
        await msg.answer("❌ Raqam!")
        return
    
    await state.update_data(pin_count=count)
    await msg.answer(f"✅ {count} ta PIN\n\nSozlamalar:", reply_markup=pin_settings_kb())

@router.callback_query(F.data.startswith("pinset_"))
async def pin_settings(cb: CallbackQuery, state: FSMContext):
    setting = cb.data.split("_")[1]
    
    # State da pin_count borligini tekshirish
    data = await state.get_data()
    if 'pin_count' not in data:
        await cb.answer("❌ Avval PIN soni kiriting!", show_alert=True)
        return
    
    if setting == "single":
        await state.update_data(multi_use=False, max_attempts=1)
        await cb.answer("✅ Bir martalik")
        await create_batch_pins(cb.message, state)
    elif setting == "multi":
        await state.update_data(multi_use=True, max_attempts=999)
        await cb.answer("✅ Ko'p martalik")
        await create_batch_pins(cb.message, state)
    elif setting == "attempts":
        await cb.message.edit_text(
            "🔢 Maksimal urinishlar sonini kiriting:\n\n"
            "Masalan: 3 (har bir PIN 3 marta ishlatiladi)"
        )
        await state.update_data(multi_use=True)
        await state.set_state(AdminStates.pin_config)

@router.message(AdminStates.pin_config)
async def pin_config(msg: Message, state: FSMContext):
    try:
        attempts = int(msg.text.strip())
        if attempts < 1:
            await msg.answer("❌ Minimum 1 urinish!")
            return
    except:
        await msg.answer("❌ Faqat raqam kiriting!")
        return
    
    await state.update_data(max_attempts=attempts)
    
    # State da barcha kerakli ma'lumotlar borligini tekshirish
    data = await state.get_data()
    required_keys = ['grade', 'topic', 'pin_count']
    missing = [k for k in required_keys if k not in data]
    
    if missing:
        await msg.answer(f"❌ Ma'lumotlar to'liq emas: {', '.join(missing)}\n\n/start dan qayta boshlang.")
        await state.clear()
        return
    
    await create_batch_pins(msg, state)

async def create_batch_pins(msg: Message, state: FSMContext):
    data = await state.get_data()
    
    # Barcha kerakli ma'lumotlarni tekshirish
    required = ['grade', 'topic', 'pin_count']
    for key in required:
        if key not in data:
            await msg.answer(f"❌ Xatolik: {key} topilmadi. Qaytadan boshlang.")
            await state.clear()
            return
    
    status = await msg.answer("⏳ PIN'lar yaratilmoqda...")
    
    pins = []
    batch_id = str(ObjectId())
    
    try:
        for i in range(data['pin_count']):
            pin = generate_pin()
            pin_doc = {
                "pin": pin,
                "batch_id": batch_id,
                "number": i + 1,
                "grade": data['grade'],
                "topic": data['topic'],
                "created_by": msg.from_user.id if hasattr(msg, 'from_user') else msg.chat.id,
                "created_at": datetime.now(),
                "expires_at": datetime.now() + timedelta(days=PIN_EXPIRY_DAYS),
                "active": True,
                "multi_use": data.get('multi_use', False),
                "max_attempts": data.get('max_attempts', 1),
                "used_count": 0,
                "used_by": [],
                "question_count": DEFAULT_QUESTION_COUNT,
                "time_limit": DEFAULT_TIME_LIMIT
            }
            await pins_col.insert_one(pin_doc)
            pins.append(pin_doc)
        
        batch_info = {
            "batch_id": batch_id,
            "grade": data['grade'],
            "topic": data['topic'],
            "pin_count": data['pin_count'],
            "question_count": DEFAULT_QUESTION_COUNT,
            "time_limit": DEFAULT_TIME_LIMIT,
            "multi_use": data.get('multi_use', False),
            "max_attempts": data.get('max_attempts', 1),
            "expiry_days": PIN_EXPIRY_DAYS,
            "created_by": msg.from_user.id if hasattr(msg, 'from_user') else msg.chat.id,
            "created_at": datetime.now()
        }
        await pin_batches_col.insert_one(batch_info)
        
        pdf_data = generate_pins_pdf(pins, batch_info)
        json_data = generate_pins_json(pins, batch_info)
        
        await status.delete()
        
        multi_text = "Ko'p martalik" if data.get('multi_use') else "Bir martalik"
        attempts_info = f"({data.get('max_attempts', 1)} urinish)" if data.get('multi_use') and data.get('max_attempts', 1) < 999 else ""
        
        await msg.answer(
            f"✅ {data['pin_count']} ta PIN yaratildi!\n\n"
            f"📚 {data['grade']}-sinf | {data['topic']}\n"
            f"🔄 {multi_text} {attempts_info}"
        )
        
        filename = data['topic'].replace('/', '-')[:30]
        await msg.answer_document(
            BufferedInputFile(pdf_data, f"PIN_{data['grade']}_{filename}.pdf"),
            caption="📄 PDF format - chop etish uchun"
        )
        await msg.answer_document(
            BufferedInputFile(json_data.encode('utf-8'), f"PIN_{data['grade']}_{filename}.json"),
            caption="💾 JSON format"
        )
        
        await state.clear()
        
    except Exception as e:
        await status.edit_text(f"❌ Xatolik: {e}")
        print(f"PIN creation error: {e}")
        await state.clear()


# ===== PIN MANAGEMENT =====
@router.message(F.text == "📋 PIN boshqaruv")
async def pin_mgmt(msg: Message):
    if msg.from_user.id not in ADMIN_IDS: return
    await msg.answer("PIN boshqaruv:", reply_markup=pin_management_kb())

@router.callback_query(F.data == "pinmgmt_list")
async def pin_list(cb: CallbackQuery):
    batches = await pin_batches_col.find().sort("created_at", -1).limit(10).to_list(10)
    if not batches:
        await cb.message.edit_text("❌ Yo'q!")
        return
    
    text = "📋 PIN to'plamlari:\n\n"
    for i, b in enumerate(batches, 1):
        used = await pins_col.count_documents({"batch_id": b['batch_id'], "used_count": {"$gt": 0}})
        text += f"{i}. {b['grade']}-sinf | {b['topic'][:20]}\n   {used}/{b['pin_count']}\n\n"
    
    btns = [[InlineKeyboardButton(text=f"{b['grade']}-sinf {b['topic'][:15]}", callback_data=f"pinbatch_{b['batch_id']}")] for b in batches]
    await cb.message.edit_text(text, reply_markup=InlineKeyboardMarkup(inline_keyboard=btns))

@router.callback_query(F.data.startswith("pinbatch_"))
async def pin_batch(cb: CallbackQuery):
    batch_id = cb.data.replace("pinbatch_", "")
    batch = await pin_batches_col.find_one({"batch_id": batch_id})
    if not batch:
        await cb.answer("❌ Topilmadi!")
        return
    
    pins = await pins_col.find({"batch_id": batch_id}).to_list(200)
    pdf_data = generate_pins_pdf(pins, batch)
    json_data = generate_pins_json(pins, batch)
    
    await cb.message.answer(f"📋 {batch['grade']}-sinf | {batch['topic']}\n{len(pins)} ta PIN")
    
    filename = batch['topic'].replace('/', '-')[:30]
    await cb.message.answer_document(BufferedInputFile(pdf_data, f"PIN_{filename}.pdf"), caption="📄 PDF")
    await cb.message.answer_document(BufferedInputFile(json_data.encode('utf-8'), f"PIN_{filename}.json"), caption="💾 JSON")

@router.callback_query(F.data == "pinmgmt_reset")
async def pin_reset(cb: CallbackQuery, state: FSMContext):
    await cb.message.edit_text("🔄 PIN kodni kiriting:")
    await state.set_state(AdminStates.pin_reset_select)

@router.message(AdminStates.pin_reset_select)
async def pin_reset_do(msg: Message, state: FSMContext):
    pin = msg.text.strip()
    pin_doc = await pins_col.find_one({"pin": pin})
    if not pin_doc:
        await msg.answer("❌ Topilmadi!")
        return
    
    await pins_col.update_one({"pin": pin}, {"$set": {"used_count": 0, "used_by": [], "active": True}})
    await msg.answer(f"✅ Reset: {pin}")
    await state.clear()

@router.callback_query(F.data == "pinmgmt_stats")
async def pin_stats(cb: CallbackQuery):
    total = await pins_col.count_documents({})
    active = await pins_col.count_documents({"active": True, "expires_at": {"$gt": datetime.now()}})
    used = await pins_col.count_documents({"used_count": {"$gt": 0}})
    await cb.message.edit_text(f"📊 Jami: {total}\nFaol: {active}\nIshlatilgan: {used}")

# ===== STUDENT TEST =====
@router.message(F.text.in_(["📝 Test boshlash", "/test"]))
async def test_start(msg: Message, state: FSMContext):
    await msg.answer("🔑 PIN kod:", reply_markup=ReplyKeyboardRemove())
    await state.set_state(StudentStates.waiting_pin)

@router.message(StudentStates.waiting_pin)
async def test_pin(msg: Message, state: FSMContext):
    pin = msg.text.strip()
    pin_data = await pins_col.find_one({"pin": pin, "active": True, "expires_at": {"$gt": datetime.now()}})
    
    if not pin_data:
        await msg.answer("❌ PIN noto'g'ri!")
        return
    
    if not pin_data.get('multi_use', False):
        if msg.from_user.id in pin_data.get('used_by', []):
            await msg.answer("❌ Siz bu testni ishlagansiz!")
            return
    else:
        attempts = pin_data.get('used_by', []).count(msg.from_user.id)
        if attempts >= pin_data.get('max_attempts', 999):
            await msg.answer(f"❌ {attempts} marta ishlagansiz!")
            return
    
    await state.update_data(pin_data=pin_data)
    await msg.answer("👤 Ism-familiya:")
    await state.set_state(StudentStates.waiting_name)

@router.message(StudentStates.waiting_name)
async def test_name(msg: Message, state: FSMContext):
    name = msg.text.strip()
    if len(name) < 3:
        await msg.answer("❌ To'liq kiriting!")
        return
    
    data = await state.get_data()
    pin_data = data['pin_data']
    
    all_q = await questions_col.find({"grade": pin_data['grade'], "topic": pin_data['topic']}).to_list(200)
    if not all_q:
        await msg.answer("❌ Savollar yo'q!")
        await state.clear()
        return
    
    count = min(pin_data.get('question_count', 10), len(all_q))
    selected = random.sample(all_q, count)
    
    for q in selected:
        if q.get('options') and isinstance(q.get('answer'), int):
            correct = q['options'][q['answer']]
            random.shuffle(q['options'])
            q['answer'] = q['options'].index(correct)
    
    # Unique ID
    test_id = str(ObjectId())
    
    session = {
        "test_id": test_id,
        "user_id": msg.from_user.id,
        "user_name": name,
        "pin": pin_data['pin'],
        "grade": pin_data['grade'],
        "topic": pin_data['topic'],
        "questions": [{
            'id': str(q['_id']),
            'text': q['text'],
            'options': q.get('options', []),
            'answer': q['answer'],
            'type': q.get('type', 'choice'),
            'images': q.get('images', [])
        } for q in selected],
        "answers": {},
        "current": 0,
        "started_at": datetime.now(),
        "time_limit": pin_data.get('time_limit', 30)
    }
    
    await state.update_data(session=session)
    await msg.answer(
        f"📝 Test: {name}\n"
        f"{pin_data['grade']}-sinf | {pin_data['topic']}\n"
        f"{count} ta savol | {session['time_limit']} daq\n\n"
        f"💡 Savollar orasida harakatlanishingiz mumkin!"
    )
    await state.set_state(StudentStates.taking_test)
    await send_question(msg, state, 0)


@router.message(F.text == "🗑 Savollarni o'chirish")
async def delete_questions_menu(msg: Message):
    if msg.from_user.id not in ADMIN_IDS:
        return
    
    total_q = await questions_col.count_documents({})
    
    await msg.answer(
        f"🗑 <b>Savollarni o'chirish</b>\n\n"
        f"📊 Bazada {total_q} ta savol bor\n\n"
        f"⚠️ O'chirilgan savollar qayta tiklanmaydi!",
        parse_mode="HTML",
        reply_markup=delete_questions_kb()
    )

@router.callback_query(F.data == "delq_by_grade")
async def delete_by_grade_start(cb: CallbackQuery, state: FSMContext):
    await cb.message.edit_text(
        "📚 Qaysi sinf savollarini o'chirmoqchisiz?",
        reply_markup=grade_kb()
    )
    await state.set_state(AdminStates.delete_by_grade)

@router.callback_query(AdminStates.delete_by_grade, F.data.startswith("grade_"))
async def delete_by_grade_confirm(cb: CallbackQuery, state: FSMContext):
    grade = int(cb.data.split("_")[1])
    count = await questions_col.count_documents({"grade": grade})
    
    if count == 0:
        await cb.message.edit_text(f"❌ {grade}-sinf uchun savollar yo'q!")
        await state.clear()
        return
    
    await state.update_data(delete_grade=grade, delete_count=count)
    
    confirm_kb = InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="✅ Ha, o'chirish", callback_data="confirm_delete_grade"),
            InlineKeyboardButton(text="❌ Bekor qilish", callback_data="cancel")
        ]
    ])
    
    await cb.message.edit_text(
        f"⚠️ <b>Tasdiqlash</b>\n\n"
        f"{grade}-sinf uchun <b>{count} ta savol</b> o'chiriladi!\n\n"
        f"Davom etamizmi?",
        parse_mode="HTML",
        reply_markup=confirm_kb
    )

@router.callback_query(F.data == "confirm_delete_grade")
async def delete_by_grade_execute(cb: CallbackQuery, state: FSMContext):
    data = await state.get_data()
    grade = data['delete_grade']
    
    result = await questions_col.delete_many({"grade": grade})
    
    await cb.message.edit_text(
        f"✅ {result.deleted_count} ta savol o'chirildi!\n"
        f"📚 Sinf: {grade}"
    )
    await state.clear()

@router.callback_query(F.data == "delq_by_topic")
async def delete_by_topic_start(cb: CallbackQuery, state: FSMContext):
    await cb.message.edit_text("📚 Avval sinfni tanlang:", reply_markup=grade_kb())
    await state.set_state(AdminStates.delete_by_topic)
    await state.update_data(delete_step="grade")

@router.callback_query(AdminStates.delete_by_topic, F.data.startswith("grade_"))
async def delete_by_topic_grade(cb: CallbackQuery, state: FSMContext):
    grade = int(cb.data.split("_")[1])
    topics = await questions_col.distinct("topic", {"grade": grade})
    
    if not topics:
        await cb.message.edit_text(f"❌ {grade}-sinf uchun mavzular yo'q!")
        await state.clear()
        return
    
    await state.update_data(delete_grade=grade)
    
    btns = []
    for topic in topics:
        count = await questions_col.count_documents({"grade": grade, "topic": topic})
        btns.append([InlineKeyboardButton(
            text=f"{topic} ({count} ta)",
            callback_data=f"deltopic_{topic}"
        )])
    
    await cb.message.edit_text(
        f"📝 {grade}-sinf\n\nMavzuni tanlang:",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=btns)
    )

@router.callback_query(AdminStates.delete_by_topic, F.data.startswith("deltopic_"))
async def delete_by_topic_confirm(cb: CallbackQuery, state: FSMContext):
    topic = cb.data.replace("deltopic_", "")
    data = await state.get_data()
    grade = data['delete_grade']
    
    count = await questions_col.count_documents({"grade": grade, "topic": topic})
    
    await state.update_data(delete_topic=topic, delete_count=count)
    
    confirm_kb = InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="✅ Ha, o'chirish", callback_data="confirm_delete_topic"),
            InlineKeyboardButton(text="❌ Yo'q", callback_data="cancel")
        ]
    ])
    
    await cb.message.edit_text(
        f"⚠️ <b>Tasdiqlash</b>\n\n"
        f"📚 {grade}-sinf\n"
        f"📝 {topic}\n"
        f"🗑 {count} ta savol o'chiriladi!\n\n"
        f"Davom etamizmi?",
        parse_mode="HTML",
        reply_markup=confirm_kb
    )

@router.callback_query(F.data == "confirm_delete_topic")
async def delete_by_topic_execute(cb: CallbackQuery, state: FSMContext):
    data = await state.get_data()
    grade = data['delete_grade']
    topic = data['delete_topic']
    
    result = await questions_col.delete_many({"grade": grade, "topic": topic})
    
    await cb.message.edit_text(
        f"✅ {result.deleted_count} ta savol o'chirildi!\n"
        f"📚 {grade}-sinf | {topic}"
    )
    await state.clear()

@router.callback_query(F.data == "delq_individual")
async def delete_individual_start(cb: CallbackQuery):
    await cb.message.edit_text("Sinfni tanlang:", reply_markup=grade_kb())
    # Bu qism oldingi "O'chirish" funksiyasiga o'xshash

@router.callback_query(F.data == "delq_all")
async def delete_all_confirm(cb: CallbackQuery, state: FSMContext):
    total = await questions_col.count_documents({})
    
    confirm_kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="⚠️ Barcha savollarni o'chirish", callback_data="confirm_delete_all")],
        [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="cancel")]
    ])
    
    await cb.message.edit_text(
        f"🚨 <b>DIQQAT!</b>\n\n"
        f"Bazadagi <b>BARCHA {total} ta savol</b> o'chiriladi!\n\n"
        f"Bu amalni bekor qilib bo'lmaydi!\n\n"
        f"Davom etamizmi?",
        parse_mode="HTML",
        reply_markup=confirm_kb
    )

@router.callback_query(F.data == "confirm_delete_all")
async def delete_all_execute(cb: CallbackQuery):
    result = await questions_col.delete_many({})
    
    await cb.message.edit_text(f"✅ Barcha {result.deleted_count} ta savol o'chirildi!")

# ===== BATAFSIL NATIJALAR =====

@router.message(F.text == "📊 Natijalar")
async def results_menu(msg: Message, state: FSMContext):
    if msg.from_user.id not in ADMIN_IDS:
        return
    
    await msg.answer(
        "📊 <b>Test natijalari</b>\n\n"
        "PIN kodni kiriting yoki:\n"
        "• <code>all</code> - barcha natijalar\n"
        "• <code>today</code> - bugungi natijalar\n"
        "• <code>week</code> - haftalik natijalar",
        parse_mode="HTML"
    )
    await state.set_state(TeacherStates.waiting_pin_for_report)

@router.message(TeacherStates.waiting_pin_for_report)
async def results_pin_entered(msg: Message, state: FSMContext):
    pin = msg.text.strip().lower()
    
    # Query yaratish
    if pin == 'all':
        query = {}
        limit = 100
    elif pin == 'today':
        today_start = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
        query = {"completed_at": {"$gte": today_start}}
        limit = 100
    elif pin == 'week':
        week_ago = datetime.now() - timedelta(days=7)
        query = {"completed_at": {"$gte": week_ago}}
        limit = 200
    else:
        query = {"pin": pin}
        limit = 100
    
    results = await results_col.find(query).sort("completed_at", -1).limit(limit).to_list(limit)
    
    if not results:
        await msg.answer("❌ Natijalar topilmadi!")
        await state.clear()
        return
    
    await state.update_data(results=results, query_type=pin)
    
    await msg.answer(
        f"✅ {len(results)} ta natija topildi!\n\n"
        f"Hisobot turini tanlang:",
        reply_markup=report_type_kb()
    )
    await state.set_state(TeacherStates.selecting_report_type)

@router.callback_query(TeacherStates.selecting_report_type, F.data.startswith("report_"))
async def generate_report(cb: CallbackQuery, state: FSMContext):
    report_type = cb.data.replace("report_", "")
    data = await state.get_data()
    results = data['results']
    query_type = data.get('query_type', 'unknown')
    
    status = await cb.message.edit_text("⏳ PDF yaratilmoqda...")
    
    try:
        if report_type == "summary":
            pdf_data = generate_summary_report(results)
            filename = f"Qisqacha_Hisobot_{query_type}_{datetime.now().strftime('%d%m%Y')}.pdf"
            caption = f"📄 Qisqacha hisobot - {len(results)} ta natija"
            
            await status.delete()
            await cb.message.answer_document(
                BufferedInputFile(pdf_data, filename),
                caption=caption
            )
        
        elif report_type == "detailed":
            pdf_data = generate_detailed_student_report(results)
            filename = f"Batafsil_Hisobot_{query_type}_{datetime.now().strftime('%d%m%Y')}.pdf"
            caption = f"📋 Batafsil hisobot - {len(results)} ta natija\n\nHar bir o'quvchi uchun alohida ma'lumotlar"
            
            await status.delete()
            await cb.message.answer_document(
                BufferedInputFile(pdf_data, filename),
                caption=caption
            )
        
        elif report_type == "both":
            # Qisqacha
            pdf_summary = generate_summary_report(results)
            await cb.message.answer_document(
                BufferedInputFile(pdf_summary, f"Qisqacha_{query_type}.pdf"),
                caption="📄 Qisqacha hisobot"
            )
            
            # Batafsil
            pdf_detailed = generate_detailed_student_report(results)
            await cb.message.answer_document(
                BufferedInputFile(pdf_detailed, f"Batafsil_{query_type}.pdf"),
                caption="📋 Batafsil hisobot"
            )
            
            await status.delete()
    
    except Exception as e:
        await status.edit_text(f"❌ Xatolik: {e}")
        print(f"PDF generation error: {e}")
    
    await state.clear()

# ===== SOZLAMALAR =====

@router.message(F.text == "⚙️ Sozlamalar")
async def settings_menu(msg: Message):
    if msg.from_user.id not in ADMIN_IDS:
        return
    
    total_q = await questions_col.count_documents({})
    total_r = await results_col.count_documents({})
    total_p = await pins_col.count_documents({})
    total_img = await images_col.count_documents({})
    
    # Xotira hisoblash
    db_stats = await db.command("dbStats")
    db_size_mb = db_stats.get('dataSize', 0) / (1024 * 1024)
    
    settings_text = f"""⚙️ <b>Tizim sozlamalari</b>

📊 <b>Ma'lumotlar bazasi:</b>
• Savollar: {total_q} ta
• Natijalar: {total_r} ta
• PIN kodlar: {total_p} ta
• Rasmlar: {total_img} ta
• Hajm: {db_size_mb:.2f} MB

🔧 <b>Konfiguratsiya:</b>
• Savol soni: {DEFAULT_QUESTION_COUNT} ta
• Vaqt: {DEFAULT_TIME_LIMIT} daqiqa
• PIN amal qilish: {PIN_EXPIRY_DAYS} kun
• Rasm hajmi: {MAX_IMAGE_SIZE//1024} KB
• Rasm o'lchami: {MAX_IMAGE_DIMENSION}px

👤 <b>Adminlar:</b> {len(ADMIN_IDS)}
"""
    
    settings_kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🔄 Bazani tozalash", callback_data="clean_db")],
        [InlineKeyboardButton(text="💾 Backup olish", callback_data="backup_db")],
        [InlineKeyboardButton(text="📊 To'liq statistika", callback_data="full_stats")]
    ])
    
    await msg.answer(settings_text, parse_mode="HTML", reply_markup=settings_kb)

@router.callback_query(F.data == "clean_db")
async def clean_database(cb: CallbackQuery):
    # Muddati o'tgan PIN'larni o'chirish
    result = await pins_col.delete_many({"expires_at": {"$lt": datetime.now()}})
    
    # Eski natijalarni o'chirish (3 oydan eski)
    three_months_ago = datetime.now() - timedelta(days=90)
    old_results = await results_col.delete_many({"completed_at": {"$lt": three_months_ago}})
    
    await cb.answer(f"✅ Tozalandi! PIN: {result.deleted_count}, Natijalar: {old_results.deleted_count}", show_alert=True)

@router.callback_query(F.data == "full_stats")
async def full_statistics(cb: CallbackQuery):
    # Sinf bo'yicha
    grades_pipeline = [
        {"$group": {"_id": "$grade", "count": {"$sum": 1}}},
        {"$sort": {"_id": 1}}
    ]
    grades = await questions_col.aggregate(grades_pipeline).to_list(10)
    
    # Qiyinlik bo'yicha
    diff_pipeline = [
        {"$group": {"_id": "$difficulty", "count": {"$sum": 1}}}
    ]
    difficulties = await questions_col.aggregate(diff_pipeline).to_list(10)
    
    # Eng yaxshi natijalar
    top_students = await results_col.find().sort("score", -1).limit(5).to_list(5)
    
    stats_text = "📊 <b>TO'LIQ STATISTIKA</b>\n\n"
    
    stats_text += "<b>Sinflar bo'yicha:</b>\n"
    for g in grades:
        stats_text += f"  {g['_id']}-sinf: {g['count']} ta\n"
    
    stats_text += "\n<b>Qiyinlik darajasi:</b>\n"
    for d in difficulties:
        stats_text += f"  {d['_id']}: {d['count']} ta\n"
    
    stats_text += "\n<b>🏆 Top 5 natijalar:</b>\n"
    for i, s in enumerate(top_students, 1):
        stats_text += f"{i}. {s['user_name']}: {s['score']}%\n"
    
    await cb.message.answer(stats_text, parse_mode="HTML")



async def send_question(msg: Message, state: FSMContext, q_index: int):
    """Savolni yuborish - YANGILANGAN"""
    data = await state.get_data()
    s = data['session']
    
    total = len(s['questions'])
    if q_index < 0 or q_index >= total:
        return
    
    # Vaqt
    elapsed = (datetime.now() - s['started_at']).total_seconds() / 60
    remain = s['time_limit'] - elapsed
    
    if remain <= 0:
        await msg.answer("⏱ Vaqt tugadi!")
        await finish_test(msg, state)
        return
    
    q = s['questions'][q_index]
    
    # Savol matni
    answered = "✅" if str(q_index) in s['answers'] else "⬜"
    text = f"{answered} <b>Savol {q_index+1}/{total}</b> | ⏱ {int(remain)} daq\n\n{q['text']}"
    
    # Rasmlarni yuborish
    if q.get('images'):
        print(f"Savol {q_index+1} da {len(q['images'])} ta rasm bor")
        for img_id in q['images']:
            img_data = await get_image(img_id)
            if img_data:
                try:
                    await msg.answer_photo(
                        BufferedInputFile(img_data, "question.jpg"),
                        caption=f"📷 Savol {q_index+1}"
                    )
                    print(f"✓ Rasm yuborildi: {img_id}")
                except Exception as e:
                    print(f"✗ Rasm yuborishda xato: {e}")
            else:
                print(f"✗ Rasm topilmadi: {img_id}")
    
    # Javob klaviaturasi
    if q['type'] == 'choice' and q.get('options'):
        # Javob berilgan bo'lsa ko'rsatish
        user_answer = s['answers'].get(str(q_index))
        if user_answer is not None:
            try:
                text += f"\n\n<i>Sizning javobingiz: {chr(65 + user_answer)}) {q['options'][user_answer][:30]}...</i>"
            except:
                pass
        
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(
                text=f"{chr(65+i)}) {opt[:40]}{'...' if len(opt)>40 else ''}",
                callback_data=f"ans_{s['test_id']}_{q_index}_{i}"
            )] for i, opt in enumerate(q['options'])
        ])
        
        await msg.answer(text, parse_mode="HTML", reply_markup=keyboard)
    else:
        await msg.answer(text + "\n\n✏️ Javobingizni yozing:", parse_mode="HTML")
        await state.set_state(StudentStates.waiting_text_answer)
    
    # Navigatsiya
    nav_kb = navigation_kb(q_index, total, s['answers'], s['test_id'])
    await msg.answer("📍 <b>Navigatsiya:</b>", parse_mode="HTML", reply_markup=nav_kb)


def generate_detailed_student_report(results: List[dict]) -> bytes:
    """Har bir o'quvchi uchun batafsil hisobot"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1.5*cm, bottomMargin=1.5*cm)
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=18,
        alignment=1,
        spaceAfter=20,
        textColor=colors.HexColor('#1a5f7a')
    )
    elements.append(Paragraph("📊 BATAFSIL TEST NATIJALARI", title_style))
    elements.append(Spacer(1, 0.5*cm))
    
    # Umumiy statistika
    total_students = len(results)
    avg_score = sum(r['score'] for r in results) / total_students if total_students > 0 else 0
    avg_time = sum(r['time_seconds'] for r in results) / total_students if total_students > 0 else 0
    
    summary_style = ParagraphStyle('Summary', parent=styles['Normal'], fontSize=10, spaceAfter=8)
    elements.append(Paragraph(f"<b>Jami o'quvchilar:</b> {total_students}", summary_style))
    elements.append(Paragraph(f"<b>O'rtacha ball:</b> {avg_score:.1f}%", summary_style))
    elements.append(Paragraph(f"<b>O'rtacha vaqt:</b> {int(avg_time//60)}:{int(avg_time%60):02d}", summary_style))
    elements.append(Paragraph(f"<b>Sana:</b> {datetime.now().strftime('%d.%m.%Y %H:%M')}", summary_style))
    elements.append(Spacer(1, 1*cm))
    
    # Har bir o'quvchi uchun
    for idx, result in enumerate(results, 1):
        # O'quvchi ma'lumotlari
        student_style = ParagraphStyle('Student', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#2c5f2d'))
        elements.append(Paragraph(f"{idx}. {result['user_name']}", student_style))
        
        info_style = ParagraphStyle('Info', parent=styles['Normal'], fontSize=9, spaceAfter=4)
        elements.append(Paragraph(f"📚 <b>Sinf:</b> {result['grade']}", info_style))
        elements.append(Paragraph(f"📝 <b>Mavzu:</b> {result['topic']}", info_style))
        elements.append(Paragraph(f"📊 <b>Ball:</b> {result['score']}% ({result['correct']}/{result['total']})", info_style))
        
        minutes = int(result['time_seconds'] // 60)
        seconds = int(result['time_seconds'] % 60)
        elements.append(Paragraph(f"⏱ <b>Vaqt:</b> {minutes}:{seconds:02d}", info_style))
        elements.append(Paragraph(f"📅 <b>Sana:</b> {result['completed_at'].strftime('%d.%m.%Y %H:%M')}", info_style))
        elements.append(Spacer(1, 0.3*cm))
        
        # Savollar jadvali
        if result.get('details'):
            table_data = [["№", "Savol", "O'quvchi javobi", "To'g'ri javob", "Natija"]]
            
            for i, detail in enumerate(result['details'], 1):
                status = "✓" if detail.get('ok') else "✗"
                status_color = colors.green if detail.get('ok') else colors.red
                
                # Javoblarni formatlash
                user_ans = detail.get('user', 'Berilmagan')
                if isinstance(user_ans, int):
                    user_ans = chr(65 + user_ans)  # 0->A, 1->B, etc
                
                correct_ans = detail.get('correct', '-')
                if isinstance(correct_ans, int):
                    correct_ans = chr(65 + correct_ans)
                
                table_data.append([
                    str(i),
                    Paragraph(detail['q'][:80] + "...", styles['Normal']) if len(detail['q']) > 80 else detail['q'],
                    str(user_ans),
                    str(correct_ans),
                    Paragraph(f"<font color='{status_color.hexval()}'>{status}</font>", styles['Normal'])
                ])
            
            # Jadval yaratish
            col_widths = [1*cm, 9*cm, 2.5*cm, 2.5*cm, 1.5*cm]
            table = Table(table_data, colWidths=col_widths)
            table.setStyle(TableStyle([
                # Header
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4a5568')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                # Body
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ]))
            
            elements.append(table)
        
        elements.append(Spacer(1, 1*cm))
        
        # Sahifa uzilishi (oxirgi o'quvchidan tashqari)
        if idx < len(results):
            elements.append(Paragraph("<para align='center'>• • •</para>", styles['Normal']))
            elements.append(Spacer(1, 0.5*cm))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer.read()


# ... (oldingi import va config bir xil) ...

# ================= PDF GENERATION - YANGILANGAN =================

def generate_detailed_student_report(results: List[dict]) -> bytes:
    """Har bir o'quvchi uchun batafsil hisobot"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1.5*cm, bottomMargin=1.5*cm)
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=18,
        alignment=1,
        spaceAfter=20,
        textColor=colors.HexColor('#1a5f7a')
    )
    elements.append(Paragraph("📊 BATAFSIL TEST NATIJALARI", title_style))
    elements.append(Spacer(1, 0.5*cm))
    
    # Umumiy statistika
    total_students = len(results)
    avg_score = sum(r['score'] for r in results) / total_students if total_students > 0 else 0
    avg_time = sum(r['time_seconds'] for r in results) / total_students if total_students > 0 else 0
    
    summary_style = ParagraphStyle('Summary', parent=styles['Normal'], fontSize=10, spaceAfter=8)
    elements.append(Paragraph(f"<b>Jami o'quvchilar:</b> {total_students}", summary_style))
    elements.append(Paragraph(f"<b>O'rtacha ball:</b> {avg_score:.1f}%", summary_style))
    elements.append(Paragraph(f"<b>O'rtacha vaqt:</b> {int(avg_time//60)}:{int(avg_time%60):02d}", summary_style))
    elements.append(Paragraph(f"<b>Sana:</b> {datetime.now().strftime('%d.%m.%Y %H:%M')}", summary_style))
    elements.append(Spacer(1, 1*cm))
    
    # Har bir o'quvchi uchun
    for idx, result in enumerate(results, 1):
        # O'quvchi ma'lumotlari
        student_style = ParagraphStyle('Student', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#2c5f2d'))
        elements.append(Paragraph(f"{idx}. {result['user_name']}", student_style))
        
        info_style = ParagraphStyle('Info', parent=styles['Normal'], fontSize=9, spaceAfter=4)
        elements.append(Paragraph(f"📚 <b>Sinf:</b> {result['grade']}", info_style))
        elements.append(Paragraph(f"📝 <b>Mavzu:</b> {result['topic']}", info_style))
        elements.append(Paragraph(f"📊 <b>Ball:</b> {result['score']}% ({result['correct']}/{result['total']})", info_style))
        
        minutes = int(result['time_seconds'] // 60)
        seconds = int(result['time_seconds'] % 60)
        elements.append(Paragraph(f"⏱ <b>Vaqt:</b> {minutes}:{seconds:02d}", info_style))
        elements.append(Paragraph(f"📅 <b>Sana:</b> {result['completed_at'].strftime('%d.%m.%Y %H:%M')}", info_style))
        elements.append(Spacer(1, 0.3*cm))
        
        # Savollar jadvali
        if result.get('details'):
            table_data = [["№", "Savol", "O'quvchi javobi", "To'g'ri javob", "Natija"]]
            
            for i, detail in enumerate(result['details'], 1):
                status = "✓" if detail.get('ok') else "✗"
                status_color = colors.green if detail.get('ok') else colors.red
                
                # Javoblarni formatlash
                user_ans = detail.get('user', 'Berilmagan')
                if isinstance(user_ans, int):
                    user_ans = chr(65 + user_ans)  # 0->A, 1->B, etc
                
                correct_ans = detail.get('correct', '-')
                if isinstance(correct_ans, int):
                    correct_ans = chr(65 + correct_ans)
                
                table_data.append([
                    str(i),
                    Paragraph(detail['q'][:80] + "...", styles['Normal']) if len(detail['q']) > 80 else detail['q'],
                    str(user_ans),
                    str(correct_ans),
                    Paragraph(f"<font color='{status_color.hexval()}'>{status}</font>", styles['Normal'])
                ])
            
            # Jadval yaratish
            col_widths = [1*cm, 9*cm, 2.5*cm, 2.5*cm, 1.5*cm]
            table = Table(table_data, colWidths=col_widths)
            table.setStyle(TableStyle([
                # Header
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4a5568')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                # Body
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ]))
            
            elements.append(table)
        
        elements.append(Spacer(1, 1*cm))
        
        # Sahifa uzilishi (oxirgi o'quvchidan tashqari)
        if idx < len(results):
            elements.append(Paragraph("<para align='center'>• • •</para>", styles['Normal']))
            elements.append(Spacer(1, 0.5*cm))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer.read()

def generate_summary_report(results: List[dict]) -> bytes:
    """Qisqacha umumiy hisobot"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=16, alignment=1, spaceAfter=20)
    elements.append(Paragraph("📊 TEST NATIJALARI - QISQACHA", title_style))
    elements.append(Spacer(1, 0.5*cm))
    
    # Jadval
    table_data = [["№", "Ism-Familiya", "Sinf", "Mavzu", "Ball", "Vaqt", "Sana"]]
    
    for i, r in enumerate(results, 1):
        minutes = int(r['time_seconds'] // 60)
        seconds = int(r['time_seconds'] % 60)
        
        # Ball rangini aniqlash
        if r['score'] >= 86:
            score_color = colors.green
        elif r['score'] >= 71:
            score_color = colors.blue
        elif r['score'] >= 56:
            score_color = colors.orange
        else:
            score_color = colors.red
        
        table_data.append([
            str(i),
            r['user_name'][:25],
            str(r['grade']),
            r['topic'][:20],
            Paragraph(f"<font color='{score_color.hexval()}'><b>{r['score']}%</b></font>", styles['Normal']),
            f"{minutes}:{seconds:02d}",
            r['completed_at'].strftime("%d.%m.%Y")
        ])
    
    table = Table(table_data, colWidths=[1*cm, 4.5*cm, 1.5*cm, 4*cm, 2*cm, 2*cm, 2.5*cm])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
    ]))
    
    elements.append(table)
    
    # Statistika
    elements.append(Spacer(1, 1*cm))
    stat_style = ParagraphStyle('Stat', parent=styles['Normal'], fontSize=10, spaceAfter=6)
    
    total = len(results)
    avg_score = sum(r['score'] for r in results) / total if total > 0 else 0
    excellent = sum(1 for r in results if r['score'] >= 86)
    good = sum(1 for r in results if 71 <= r['score'] < 86)
    satisfactory = sum(1 for r in results if 56 <= r['score'] < 71)
    poor = sum(1 for r in results if r['score'] < 56)
    
    elements.append(Paragraph("<b>Statistika:</b>", stat_style))
    elements.append(Paragraph(f"Jami: {total} | O'rtacha: {avg_score:.1f}%", stat_style))
    elements.append(Paragraph(f"A'lo (86-100%): {excellent} | Yaxshi (71-85%): {good}", stat_style))
    elements.append(Paragraph(f"Qoniqarli (56-70%): {satisfactory} | Qoniqarsiz (<56%): {poor}", stat_style))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer.read()




@router.callback_query(F.data.startswith("ans_"))
async def answer_selected(cb: CallbackQuery, state: FSMContext):
    """Javob tanlash"""
    parts = cb.data.split("_")
    test_id = parts[1]
    q_index = int(parts[2])
    answer = int(parts[3])
    
    data = await state.get_data()
    s = data.get('session', {})
    
    if s.get('test_id') != test_id:
        await cb.answer("❌ Eski test!")
        return
    
    # Javobni saqlash
    s['answers'][str(q_index)] = answer
    await state.update_data(session=s)
    
    await cb.answer(f"✅ Javob saqlandi: {chr(65 + answer)}")

@router.callback_query(F.data.startswith("nav_"))
async def navigate(cb: CallbackQuery, state: FSMContext):
    """Navigatsiya - TUZATILGAN"""
    parts = cb.data.split("_")
    
    # nav_info uchun
    if len(parts) < 3:
        await cb.answer("ℹ️ Savollarni tanlang yoki oldingi/keyingi tugmalaridan foydalaning")
        return
    
    test_id = parts[1]
    direction = parts[2]
    
    data = await state.get_data()
    s = data.get('session', {})
    
    if s.get('test_id') != test_id:
        await cb.answer("❌ Eski test!")
        return
    
    current = s.get('current', 0)
    total = len(s['questions'])
    
    if direction == "prev":
        new_index = max(0, current - 1)
    elif direction == "next":
        new_index = min(total - 1, current + 1)
    else:
        await cb.answer("ℹ️ Navigatsiya")
        return
    
    s['current'] = new_index
    await state.update_data(session=s)
    
    try:
        await cb.message.delete()
    except:
        pass
    
    await send_question(cb.message, state, new_index)

@router.callback_query(F.data.startswith("goto_"))
async def goto_question(cb: CallbackQuery, state: FSMContext):
    """Savolga o'tish - TUZATILGAN"""
    parts = cb.data.split("_")
    
    if len(parts) < 3:
        await cb.answer("❌ Xato!")
        return
    
    test_id = parts[1]
    q_index = int(parts[2])
    
    data = await state.get_data()
    s = data.get('session', {})
    
    if s.get('test_id') != test_id:
        await cb.answer("❌ Eski test!")
        return
    
    s['current'] = q_index
    await state.update_data(session=s)
    
    try:
        await cb.message.delete()
    except:
        pass
    
    await send_question(cb.message, state, q_index)

@router.callback_query(F.data.startswith("ans_"))
async def answer_selected(cb: CallbackQuery, state: FSMContext):
    """Javob tanlash - TUZATILGAN"""
    parts = cb.data.split("_")
    
    if len(parts) < 4:
        await cb.answer("❌ Xato!")
        return
    
    test_id = parts[1]
    q_index = int(parts[2])
    answer = int(parts[3])
    
    data = await state.get_data()
    s = data.get('session', {})
    
    if s.get('test_id') != test_id:
        await cb.answer("❌ Eski test!")
        return
    
    # Javobni saqlash
    s['answers'][str(q_index)] = answer
    await state.update_data(session=s)
    
    try:
        await cb.message.edit_text(
            cb.message.text + f"\n\n✅ <i>Javob saqlandi: {chr(65 + answer)}</i>",
            parse_mode="HTML"
        )
    except:
        pass
    
    await cb.answer(f"✅ Javob: {chr(65 + answer)}")

@router.callback_query(F.data.startswith("finish_"))
async def finish_confirm(cb: CallbackQuery, state: FSMContext):
    """Testni yakunlash - TUZATILGAN"""
    parts = cb.data.split("_")
    
    if len(parts) < 2:
        await cb.answer("❌ Xato!")
        return
    
    test_id = parts[1]
    
    data = await state.get_data()
    s = data.get('session', {})
    
    if s.get('test_id') != test_id:
        await cb.answer("❌ Eski test!")
        return
    
    answered = len(s.get('answers', {}))
    total = len(s.get('questions', []))
    
    if answered < total:
        unanswered = total - answered
        confirm_text = (
            f"⚠️ {unanswered} ta savolga javob berilmagan!\n\n"
            f"Javob berilmagan savollarga 0 ball beriladi.\n\n"
            f"Yakunlashni xohlaysizmi?"
        )
        await cb.answer(confirm_text, show_alert=True)
        
        # Tasdiqlash klaviaturasi
        kb = InlineKeyboardMarkup(inline_keyboard=[
            [
                InlineKeyboardButton(text="✅ Ha, yakunla", callback_data=f"finishyes_{test_id}"),
                InlineKeyboardButton(text="❌ Yo'q", callback_data=f"finishno_{test_id}")
            ]
        ])
        await cb.message.edit_reply_markup(reply_markup=kb)
    else:
        await finish_test(cb.message, state)

@router.callback_query(F.data.startswith("finishyes_"))
async def finish_yes(cb: CallbackQuery, state: FSMContext):
    await cb.answer("Test yakunlanmoqda...")
    await finish_test(cb.message, state)

@router.callback_query(F.data.startswith("finishno_"))
async def finish_no(cb: CallbackQuery, state: FSMContext):
    await cb.answer("Testni davom ettiring")
    data = await state.get_data()
    s = data.get('session', {})
    
    # Navigatsiya klaviaturasini qaytarish
    nav_kb = navigation_kb(s.get('current', 0), len(s['questions']), s['answers'], s['test_id'])
    await cb.message.edit_reply_markup(reply_markup=nav_kb)


async def finish_test(msg: Message, state: FSMContext):
    """Testni yakunlash"""
    data = await state.get_data()
    s = data['session']
    
    correct = sum(1 for i, q in enumerate(s['questions']) 
                  if s['answers'].get(str(i)) == q['answer'])
    
    total = len(s['questions'])
    time_sec = (datetime.now() - s['started_at']).total_seconds()
    score = round(correct / total * 100, 1)
    
    await results_col.insert_one({
        "user_id": s['user_id'],
        "user_name": s['user_name'],
        "pin": s['pin'],
        "grade": s['grade'],
        "topic": s['topic'],
        "score": score,
        "correct": correct,
        "total": total,
        "time_seconds": time_sec,
        "completed_at": datetime.now()
    })
    
    await pins_col.update_one(
        {"pin": s['pin']},
        {"$push": {"used_by": s['user_id']}, "$inc": {"used_count": 1}}
    )
    
    m, sec = int(time_sec // 60), int(time_sec % 60)
    emoji = "🏆" if score >= 86 else "🥈" if score >= 71 else "🥉" if score >= 56 else "📝"
    
    await msg.answer(
        f"{emoji} <b>Test yakunlandi!</b>\n\n"
        f"👤 {s['user_name']}\n"
        f"📚 {s['grade']}-sinf | {s['topic']}\n\n"
        f"✅ To'g'ri: {correct}/{total}\n"
        f"📊 Ball: {score}%\n"
        f"⏱ Vaqt: {m}:{sec:02d}",
        parse_mode="HTML",
        reply_markup=student_menu()
    )
    
    # Admin xabari
    pin_data = data.get('pin_data', {})
    if admin_id := pin_data.get('created_by'):
        try:
            await bot.send_message(
                admin_id,
                f"📊 Yangi natija!\n\n{s['user_name']}: {score}% ({correct}/{total})"
            )
        except:
            pass
    
    await state.clear()

# ===== RESULTS & STATS =====
@router.message(F.text == "📊 Natijalar")
async def results_menu(msg: Message, state: FSMContext):
    if msg.from_user.id not in ADMIN_IDS: return
    await msg.answer("PIN (yoki 'all'):")
    await state.set_state(TeacherStates.waiting_pin_for_report)

@router.message(TeacherStates.waiting_pin_for_report)
async def results_gen(msg: Message, state: FSMContext):
    pin = msg.text.strip()
    query = {} if pin.lower() == 'all' else {"pin": pin}
    res = await results_col.find(query).sort("completed_at", -1).limit(50).to_list(50)
    
    if not res:
        await msg.answer("❌ Yo'q!")
        await state.clear()
        return
    
    text = "📊 <b>Natijalar:</b>\n\n"
    for i, r in enumerate(res, 1):
        text += f"{i}. {r['user_name']}: {r['score']}% ({r['correct']}/{r['total']})\n"
    
    await msg.answer(text, parse_mode="HTML")
    await state.clear()

@router.message(F.text == "📈 Statistika")
async def stats(msg: Message):
    if msg.from_user.id not in ADMIN_IDS: return
    q = await questions_col.count_documents({})
    r = await results_col.count_documents({})
    p = await pins_col.count_documents({"active": True})
    imgs = await images_col.count_documents({})
    await msg.answer(f"📈 Savollar: {q}\nTestlar: {r}\nPIN: {p}\n🖼 Rasmlar: {imgs}")

@router.message(F.text == "📊 Natijalarim")
async def my_res(msg: Message):
    res = await results_col.find({"user_id": msg.from_user.id}).sort("completed_at", -1).limit(10).to_list(10)
    if not res:
        await msg.answer("Yo'q!")
        return
    text = "📊 Natijalaringiz:\n\n"
    for i, r in enumerate(res, 1):
        text += f"{i}. {r['topic']}: {r['score']}%\n"
    await msg.answer(text)

# ===== MAIN =====
async def main():
    print("🔄 MongoDB ga ulanmoqda...")
    
    for attempt in range(3):
        try:
            await client.admin.command('ping')
            print("✅ MongoDB ulandi!")
            break
        except Exception as e:
            print(f"⚠️ Urinish {attempt + 1}/3")
            if attempt == 2:
                print("❌ Ulanmadi!")
                return
            await asyncio.sleep(5)
    
    try:
        await questions_col.create_index([("grade", 1), ("topic", 1)])
        await results_col.create_index([("user_id", 1), ("completed_at", -1)])
        await results_col.create_index("completed_at")  # YANGI
        await pins_col.create_index("pin", unique=True)
        await pins_col.create_index("expires_at")  # YANGI
        await images_col.create_index("hash", unique=True)
        print("✅ Indexlar yaratildi!")
    except Exception as e:
        print(f"⚠️ Index: {e}")
    
    dp.include_router(router)
    
    # Statistika
    total_q = await questions_col.count_documents({})
    total_r = await results_col.count_documents({})
    total_img = await images_col.count_documents({})
    
    print(f"🚀 Bot ishga tushdi!")
    print(f"👤 Adminlar: {ADMIN_IDS}")
    print(f"📊 Savollar: {total_q}")
    print(f"📝 Natijalar: {total_r}")
    print(f"🖼 Rasmlar: {total_img}")
    print("\n✅ Bot ishlayapti...\n")
    
    await dp.start_polling(bot)

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\n👋 Bot to'xtatildi!")